#!/usr/bin/env python3
"""
Test script for document processing functionality
Tests Word document and image processing capabilities
"""

import requests
import os
import tempfile

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    print("Warning: Pillow not available")

try:
    import pytesseract
    PYTESSERACT_AVAILABLE = True
except ImportError:
    PYTESSERACT_AVAILABLE = False
    print("Warning: pytesseract not available")

def create_test_word_document():
    """Create a test Word document"""
    if not DOCX_AVAILABLE:
        print("Error: python-docx not available, cannot create Word document")
        return None
    
    try:
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test document for the RAG system.')
        doc.add_paragraph('It contains sample text that should be processed.')
        doc.add_paragraph('The system should be able to extract this text.')
        
        # Save to temporary file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            return tmp.name
    except Exception as e:
        print(f"Error creating Word document: {e}")
        return None

def create_test_image_with_text():
    """Create a simple test image with text (would need actual image creation in real scenario)"""
    # For testing purposes, we'll just return a path to an existing image
    # In a real test, you'd create an image with text
    return None

def test_word_document_processing():
    """Test Word document processing functionality"""
    print("Testing Word document processing...", flush=True)
    
    if not DOCX_AVAILABLE:
        print("Warning: python-docx not available, skipping Word document test", flush=True)
        return
    
    # Create test document
    doc_path = create_test_word_document()
    if not doc_path:
        print("Failed to create test document", flush=True)
        return
    
    try:
        # Test direct processing
        doc = Document(doc_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        print(f"Extracted text: {text}", flush=True)
        
        # Test API endpoint (if server is running)
        try:
            with open(doc_path, 'rb') as f:
                files = {'file': f}
                api_url = os.getenv('API_URL', 'http://localhost:5000')
                response = requests.post(f'{api_url}/upload', files=files)
                
            if response.status_code == 200:
                result = response.json()
                print(f"API Response: {result}")
            else:
                print(f"API Error: {response.status_code} - {response.text}")
        except requests.exceptions.ConnectionError:
            print("API server not running - skipping API test")
    
    finally:
        # Clean up
        if os.path.exists(doc_path):
            os.remove(doc_path)

def test_ocr_functionality():
    """Test OCR functionality"""
    print("Testing OCR functionality...", flush=True)
    
    if not PYTESSERACT_AVAILABLE:
        print("Warning: pytesseract not available, skipping OCR test", flush=True)
        return
    
    if not PIL_AVAILABLE:
        print("Warning: Pillow not available, skipping OCR test", flush=True)
        return
    
    # Test with a simple image (you'd need an actual image with text)
    # This is a basic test to ensure imports work
    try:
        # Test OCR import and basic functionality
        print("OCR functionality available", flush=True)
        try:
            version = pytesseract.get_tesseract_version()
            print(f"Tesseract version info: {version}", flush=True)
        except Exception as e:
            print(f"Note: Tesseract engine not installed - {e}", flush=True)
            print("OCR functionality requires Tesseract to be installed separately", flush=True)
            print("Install from: https://github.com/UB-Mannheim/tesseract/wiki", flush=True)
    except Exception as e:
        print(f"OCR test failed: {e}", flush=True)

if __name__ == "__main__":
    print("Starting document processing tests...", flush=True)
    print(f"Available libraries:", flush=True)
    print(f"  - python-docx: {'✓' if DOCX_AVAILABLE else '✗'}", flush=True)
    print(f"  - Pillow: {'✓' if PIL_AVAILABLE else '✗'}", flush=True)
    print(f"  - pytesseract: {'✓' if PYTESSERACT_AVAILABLE else '✗'}", flush=True)
    print(flush=True)
    
    # Test Word document processing
    test_word_document_processing()
    
    # Test OCR functionality
    test_ocr_functionality()
    
    print("\nTests completed!", flush=True)